#!/usr/bin/env python3
"""
Script to read Word documents from Resources folder
and extract their content for context understanding
"""

import os
from docx import Document

def read_docx(file_path):
    """Read a Word document and return its text content"""
    doc = Document(file_path)

    full_text = []

    # Read all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Read all tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\n".join(full_text)

def main():
    resources_dir = r"C:\Users\DELL\Desktop\sanctuari\Resources"

    files_to_read = [
        "component-library-and-user-stories.docx",
        "initial-prompt.docx",
        "technical-specifications.docx"
    ]

    print("=" * 80)
    print("SANCTUARI PROJECT DOCUMENTATION")
    print("=" * 80)
    print()

    for filename in files_to_read:
        file_path = os.path.join(resources_dir, filename)

        if os.path.exists(file_path):
            print(f"\n{'=' * 80}")
            print(f"FILE: {filename}")
            print(f"{'=' * 80}\n")

            try:
                content = read_docx(file_path)
                print(content)
                print()
            except Exception as e:
                print(f"Error reading {filename}: {e}")
        else:
            print(f"\n[WARNING] File not found: {file_path}")

    print("\n" + "=" * 80)
    print("END OF DOCUMENTATION")
    print("=" * 80)

if __name__ == "__main__":
    main()
